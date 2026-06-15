from docx import Document

def parse_docx(path):
    doc = Document(path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        t = []
        for row in table.rows:
            t.append([cell.text for cell in row.cells])
        tables.append(t)

    return {
        "paragraphs": paragraphs,
        "tables": tables
    }